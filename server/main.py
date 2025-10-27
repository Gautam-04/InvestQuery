from flask import Flask, jsonify, request
from flask_cors import CORS
import requests
import spacy
import re
import json
import google.generativeai as genai
import os
from dotenv import load_dotenv
from alpha_vantage.timeseries import TimeSeries
from alpha_vantage.fundamentaldata import FundamentalData
import yfinance as yf
from PyPDF2 import PdfReader
import pandas as pd
import os
import tempfile
from docx import Document
from pytrends.request import TrendReq
from rag_utils import build_index_from_text, retrieve_relevant_chunks
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document


load_dotenv()


app = Flask(__name__)
CORS(app)

nlp = spacy.load("en_core_web_sm")

os.environ["KMP_DUPLICATE_LIB_OK"]="TRUE"
genai.configure(api_key=os.environ['GEMINI_API_KEY'])
MODEL = genai.GenerativeModel('gemini-2.5-flash')

ALPHA_VANTAGE_KEY = os.getenv("ALPHA_VANTAGE_KEY_2")

ts = TimeSeries(key=ALPHA_VANTAGE_KEY, output_format='json')
fd = FundamentalData(key=ALPHA_VANTAGE_KEY, output_format='json')

def extract_text_from_pdf(path):
    reader = PdfReader(path)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

@app.route("/analyze", methods=["POST"])
def analyze():
    question = request.form.get("question", "")

    # Case 1: File uploaded → Document-based QA
    if "file" in request.files:
        file = request.files["file"]
        with tempfile.NamedTemporaryFile(delete=False) as temp:
            file.save(temp.name)
            text = extract_text_from_pdf(temp.name)

        # Split and embed
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        docs = [Document(page_content=t) for t in splitter.split_text(text)]
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(docs, embeddings)
        retriever = vectorstore.as_retriever()

        # Get top context
        related_docs = retriever.get_relevant_documents(question or "summarize financial performance")
        context = "\n".join([d.page_content for d in related_docs])

        prompt = f"""
        You are an expert financial analyst AI assistant.

The following is a financial document (could be a report, statement, or dataset).
You will answer the user's question based **strictly** based on the document's content.
If any data is missing, mention that clearly without guessing.

---
### 🧾 Document Content:
{context}

---
### ❓ User Question:
{question}

---
### 🎯 Instructions:
- Provide a **comprehensive, structured, and well-formatted** answer in **Markdown**.
- Include:
  - Relevant **tables** (if numeric data is available)
  - **Comparative analysis** between years, quarters, or categories
  - **Percentage changes, growth trends, and ratios** when applicable
  - **Color-coded indicators** (using emojis like 🟢 for positive growth, 🔴 for decline, 🟡 for neutral)
  - **Key financial metrics** (Revenue, Profit, Expenses, Margins, etc.)
  - **Actionable insights** in bullet points
  - Please format the response using markdown syntax for better readability. If adding table try to have better formatting. the current formating is bad.
- End with a **short conclusion** summarizing the key takeaways.
"""
        answer = MODEL.generate_content(prompt)
        return jsonify({"answer": answer.text})

    # Case 2: No file → Financial MCQ or Recommendation Generation
    elif question:
        prompt = f"""
You are an expert AI Financial Analyst.

Your task is to analyze the following query and generate a structured professional response.

Query: "{question}"

### 🔹 Related Analytical MCQs
Generate **3 high-quality multiple-choice questions** related to this topic for learning or assessment purposes.  
Each question must have:
- 1 correct answer  
- 3 plausible distractors  

Use **clear markdown formatting**:
- Each question in **bold**  
- Each option (A, B, C, D) must appear on a **new line**  and bold the options (A, B, C, D) as well.
- Add one blank line between questions  
- Highlight the correct answer with a ✅ emoji at the end  

**Output Format Example:**

**Q1. What is the primary purpose of creating a personal budget?**  
A. To increase your investment returns immediately.  
B. To track your income and expenses to manage your money effectively.  
C. To eliminate all your debts instantly.  
D. To predict future stock market movements.  
✅ **Correct Answer:** B  

Ensure the tone is **formal**, **data-driven**, and clearly formatted for web display.
"""


        answer = MODEL.generate_content(prompt)
        return jsonify({"answer": answer.text})

    return jsonify({"answer": "Please upload a document or provide a question."})


# Extract money (like "10,000 rupees" / "$500")
def extract_amount(text):
    doc = nlp(text)
    money = []
    for ent in doc.ents:
        if ent.label_ in ["MONEY", "QUANTITY", "CARDINAL"]:
            money.append(ent.text)
    # fallback regex (like "10000 rs")
    regex_match = re.findall(r"\d+(?:,\d+)?\s?(?:rs|rupees|dollars|usd|inr)?", text.lower())
    money.extend(regex_match)
    return money if money else ["Not specified"]

# Detect risk (low, medium, high)
def extract_risk(text):
    risks = []
    if re.search(r"\blow risk\b|\blow\b", text.lower()):
        risks.append("Low risk")
    elif re.search(r"\bmedium risk\b|\bmoderate\b|\baverage\b", text.lower()):
        risks.append("Medium risk")
    elif re.search(r"\bhigh risk\b|\brisky\b|\bvolatile\b", text.lower()):
        risks.append("High risk")
    return risks if risks else ["Not specified"]

# Extract duration (like "2 years", "6 months")
def extract_duration(text):
    regex_match = re.findall(r"\d+\s?(?:year|years|month|months|days|day)", text.lower())
    return regex_match if regex_match else ["Not specified"]

# Extract investment domain keywords
def extract_domain(text):
    keywords = ["stocks", "mutual funds", "crypto", "real estate", "bonds", "gold", "etf", "etfs", "commodities", "savings account", "fixed deposits", "mfs"]
    found = [word for word in keywords if word in text.lower()]
    return found if found else ["Not specified"]

def preferred_stock_type(text):
    keywords = ["Large Cap", "Mid Cap", "Small Cap", "Blue Chip", "Growth", "Value"]
    found = [word for word in keywords if word.lower() in text.lower()]
    return found if found else ["Not specified"]

def preferred_country(text):
    keywords = ["USA", "India", "UK", "Germany", "China", "Japan", "Canada", "Australia", "US", "UAE"]
    found = [word for word in keywords if word.lower() in text.lower()]
    return found if found else ["Not specified"]

def industry_type(text):
    keywords = ["Technology", "Healthcare", "Finance", "Energy", "Consumer Goods", "Utilities", "Real Estate", "Telecommunications", "Industrials", "Materials", "Banks", "Insurance", "Pharmaceuticals", "Automotive", "Retail", "Pharma", "Banking"]
    found = [word for word in keywords if word.lower() in text.lower()]
    return found if found else ["Not specified"]

def generate_response(prompt):
    response = MODEL.generate_content(f'''{prompt}''')
    return response.text

def get_country_ticker(stock_symbol, country):
    country = country.strip().lower()
    exchange_mapping = {
        "india": ".NS",      
        "nse": ".NS",
        "uk": ".L",           
        "london": ".L",
        "germany": ".DE",     
        "china shanghai": ".SS",        
        "china shenzhen": ".SZ",
        "japan": ".T",         
        "canada": ".TO",      
        "australia": ".AX", 
        "uae": ".DFM"     
    }
    suffix = exchange_mapping.get(country, "")  
    return stock_symbol + suffix

# =========================================================
# ------------------ Trends Based Search ------------------
# =========================================================

@app.route('/trends', methods=['GET'])
def get_trends():
    try:
        url = f"https://www.alphavantage.co/query?function=SECTOR&apikey={ALPHA_VANTAGE_KEY}"
        r = requests.get(url)
        data = r.json()
        return jsonify(data), 200
    except Exception as e:
        print("Error fetching trends:", e)
        return jsonify({"error": "Failed to fetch trends"}), 500

@app.route('/api/nlpinput', methods=['POST'])
def nlpinput():
    data = request.get_json()
    searchInput = data.get('search', "")

    structured_data = []

    # Extract keywords
    extracted = {
        "amount": extract_amount(searchInput),
        "risk": extract_risk(searchInput),
        "duration": extract_duration(searchInput),
        "domain": extract_domain(searchInput),
        "preferred_stock_type": preferred_stock_type(searchInput),
        "preferred_country": preferred_country(searchInput),
        "industry_type": industry_type(searchInput)
    }

    # Refined prompt for AI
    prompt = f"""
    Based on the following investment preferences, suggest suitable investments. Only provide the names of stocks, ETFs, mutual funds, or crypto coins that match the criteria. 
    Return the final output strictly as a JSON array of strings, without any explanations, strategies, or extra text.

    Investment Preferences:
    Amount: {', '.join(extracted['amount'])}
    Risk Tolerance: {', '.join(extracted.get('risk', ['Not specified']))}
    Duration: {', '.join(extracted.get('duration', ['Not specified']))}
    Domain: {', '.join(extracted.get('domain', ['Not specified']))}
    Preferred Stock Type: {', '.join(extracted.get('preferred_stock_type', ['Not specified']))}
    Preferred Country: {', '.join(extracted.get('preferred_country', ['Not specified']))}
    Industry Type: {', '.join(extracted.get('industry_type', ['Not specified']))}

    Example output: ["HDFCBANK", "Lodha", "AAPL"] return me the tickers name only not full names.
    """

    # Call AI model
    recommendation = generate_response(prompt)

    # Clean AI response and parse JSON array
    cleanedRecommendations = re.sub(r'```json|```', '', recommendation).strip()
    try:
        stocks_list = json.loads(cleanedRecommendations)
    except json.JSONDecodeError:
        stocks_list = []
        print("Failed to parse AI recommendation:", recommendation)

    # Map stock symbols to tickers based on country
    countries = extracted.get('preferred_country', ["Not specified"])
    tickers = [get_country_ticker(stock, countries[0]) for stock in stocks_list]
    print(tickers)

    # Fetch stock data from Alpha Vantage
    for ticker in tickers:
        print(f"Fetching data for {yf.Ticker(ticker)}...")
        try:
            country = countries[0].lower()
            yf_ticker = yf.Ticker(ticker)
            info = yf_ticker.info
            # if "india" in country:
            #     # 🇮🇳 Use yfinance
            #     yf_ticker = yf.Ticker(ticker)
            #     info = yf_ticker.info

            # else:
            #     # 🇺🇸 or others: use Alpha Vantage
            #     url = f"https://www.alphavantage.co/query?function=OVERVIEW&symbol={ticker}&apikey={ALPHA_VANTAGE_KEY}"
            #     r = requests.get(url)
            #     fundamentals = r.json()

            #     stock_json = {
            #         "symbol": fundamentals.get("Symbol", ticker),
            #         "name": fundamentals.get("Name", ticker),
            #         "exchange": ticker.split('.')[-1] if '.' in ticker else "US",
            #         "sector": fundamentals.get("Sector", "N/A"),
            #         "current_price": {
            #             "value": fundamentals.get("50DayMovingAverage", "N/A"),
            #             "currency": "USD",
            #         },
            #         "details": {
            #             "market_cap": fundamentals.get("MarketCapitalization", "N/A"),
            #             "pe_ratio": fundamentals.get("PERatio", "N/A"),
            #             "eps": fundamentals.get("EPS", "N/A"),
            #             "ebitda": fundamentals.get("EBITDA", "N/A"),
            #             "roe": fundamentals.get("ReturnOnEquityTTM", "N/A"),
            #             "dividend_yield": fundamentals.get("DividendYield", "N/A"),
            #         }
            #     }

            structured_data.append(info)

        except Exception as e:
            print(f"Error fetching {ticker}: {e}")

    # Combine all extracted words
    combined = extracted["amount"] + extracted["risk"] + extracted["domain"] + extracted["duration"] + extracted["preferred_stock_type"] + extracted["preferred_country"] + extracted["industry_type"]

    return jsonify({
        "extracted_words": combined,
        "structured_data": structured_data,
        "ai_raw_response": recommendation
    }), 200


# =========================================================
# ------------------ FILE UPLOAD HANDLER ------------------
# =========================================================
def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        reader = PdfReader(file_path)
        text = " ".join(page.extract_text() for page in reader.pages if page.extract_text())
        return text

    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    elif ext in [".csv", ".xlsx"]:
        df = pd.read_excel(file_path) if ext == ".xlsx" else pd.read_csv(file_path)
        return df.to_string(index=False)

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    else:
        return "Unsupported file format."

def analyze_document(content, question=None):
    if question:
        prompt = f"""
You are an expert financial analyst AI assistant.

The following is a financial document (could be a report, statement, or dataset).
You will answer the user's question based **strictly** based on the document's content.
If any data is missing, mention that clearly without guessing.

---
### 🧾 Document Content:
{content}

---
### ❓ User Question:
{question}

---
### 🎯 Instructions:
- Provide a **comprehensive, structured, and well-formatted** answer in **Markdown**.
- Include:
  - Relevant **tables** (if numeric data is available)
  - **Comparative analysis** between years, quarters, or categories
  - **Percentage changes, growth trends, and ratios** when applicable
  - **Color-coded indicators** (using emojis like 🟢 for positive growth, 🔴 for decline, 🟡 for neutral)
  - **Key financial metrics** (Revenue, Profit, Expenses, Margins, etc.)
  - **Actionable insights** in bullet points
- End with a **short conclusion** summarizing the key takeaways.
"""
    else:
        prompt = """
You are a senior financial report analyst AI.
You are given one or more financial or business documents.

Your task is to generate a **comprehensive and visually rich Markdown summary** that helps a reader quickly understand the key financial story.

---
### 📄 Document Content:
{content}

---
### 🧠 Your Output Must Include:
1. **Executive Summary** — 3–5 sentences highlighting the overall performance.
2. **Financial Highlights Table** — Include metrics like Revenue, Profit, Expenses, EPS, YoY Growth, etc.  
   If the data is missing, infer the structure from context.
3. **Growth & Trend Analysis** — Identify:
   - Positive trends 🟢  
   - Declines 🔴  
   - Stable areas 🟡  
   Use arrows (⬆️ ⬇️ ➡️) to represent movements.
4. **Sentiment & Tone Summary** — Determine if the report tone is optimistic, neutral, or negative.
5. **Risk & Outlook Section** — Mention key risks, warnings, or opportunities.
6. **Graphical Hints (textual)** — Represent bar or line graph trends using simple ASCII bars or emoji (example: 📈📉).
7. **Well-Spaced Markdown Layout** — Use proper headings, bold, bullet points, and spacing for readability.

---
### ⚙️ Formatting Rules:
- Use Markdown headings (##, ###)
- Use tables for numerical data
- Keep tone analytical and professional
- Avoid generic filler text — always tie insights back to the document
- If document is large, summarize each section individually before concluding
"""

    response = MODEL.generate_content(prompt)
    return response.text.strip()


@app.route("/api/uploadDoc", methods=["POST"])
def uploadDoc():
    try:
        # Get uploaded file
        file = request.files['file']
        question = request.form.get('question', '')

        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp:
            temp_path = temp.name
            file.save(temp_path)

        # Close the file before re-opening (important for Windows!)
        extracted_text = extract_text_from_file(temp_path)

        # Combine question + document text
        prompt = f"""
        You are a financial analyst AI.
        The following document contains financial data or reports.

        If the user has asked a question, answer it based on the document.
        If not, summarize the key financial highlights, performance, and trends.

        Question: {question}

        Document content:
        {extracted_text[:4000]}  # limit to avoid token overflow
        """

        response = MODEL.generate_content(prompt)
        ai_text = response.text

        # Clean up temporary file
        os.remove(temp_path)

        return jsonify({"result": ai_text})

    except Exception as e:
        print("Error:", e)
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
